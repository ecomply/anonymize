from flask import Flask, request, jsonify, send_file, redirect, make_response
from flask_cors import CORS
import os
import tempfile
import zipfile
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from docx import Document
import fitz  # PyMuPDF
from flasgger import Swagger
import logging
from transformers import pipeline
import io

app = Flask(__name__, static_folder='static')
CORS(app)  # Enable CORS for all routes

# Simplified Swagger configuration
swagger_config = {
    "headers": [],
    "specs": [
        {
            "endpoint": "apispec",
            "route": "/apispec.json",
            "rule_filter": lambda rule: True,  # all in
            "model_filter": lambda tag: True,  # all in
        }
    ],
    "static_url_path": "/flasgger_static",
    "swagger_ui": True,
    "specs_route": "/api-docs"
}

swagger = Swagger(app, config=swagger_config)

# Initialize Presidio engines
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

# Initialize the transformers pipeline
transformer_anonymizer = pipeline("text-generation", model="distilgpt2")

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    text = ""
    with fitz.open(file_path) as pdf:
        for page in pdf:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    doc = Document(file_path)
    return "\\\\n".join([paragraph.text for paragraph in doc.paragraphs])

def anonymize_text_presidio(text):
    """Anonymize text using Presidio."""
    results = analyzer.analyze(text=text, entities=[], language="en")
    anonymized_text = anonymizer.anonymize(text=text, analyzer_results=results).text
    return anonymized_text

def anonymize_text_transformer(text):
    """Anonymize text using the transformers pipeline with a fallback mechanism."""
    try:
        # Generate anonymized text
        response = transformer_anonymizer(text, max_length=1024, do_sample=True, temperature=0.7)
        
        # Extract the generated text from the response
        if isinstance(response, list) and len(response) > 0:
            anonymized_text = response[0]['generated_text']
        else:
            anonymized_text = str(response)
    except Exception as e:
        logging.error(f"Transformer anonymization failed: {str(e)}. Falling back to Presidio.")
        anonymized_text = anonymize_text_presidio(text)
    
    return anonymized_text

def anonymize_pdf_presidio(input_path, output_path):
    """Anonymize a PDF file using Presidio."""
    with fitz.open(input_path) as pdf:
        doc = fitz.open()
        for page in pdf:
            text = page.get_text()
            anonymized_text = anonymize_text_presidio(text)
            new_page = doc.new_page(width=page.rect.width, height=page.rect.height)
            new_page.insert_text((72, 72), anonymized_text)
        doc.save(output_path)

def anonymize_pdf_transformer(input_path, output_path):
    """Anonymize a PDF file using transformer model."""
    with fitz.open(input_path) as pdf:
        doc = fitz.open()
        for page in pdf:
            text = page.get_text()
            anonymized_text = anonymize_text_transformer(text)
            new_page = doc.new_page(width=page.rect.width, height=page.rect.height)
            new_page.insert_text((72, 72), anonymized_text)
        doc.save(output_path)

def anonymize_docx_presidio(input_path, output_path):
    """Anonymize a DOCX file using Presidio."""
    doc = Document(input_path)
    for para in doc.paragraphs:
        para.text = anonymize_text_presidio(para.text)
    doc.save(output_path)

def anonymize_docx_transformer(input_path, output_path):
    """Anonymize a DOCX file using transformer model."""
    doc = Document(input_path)
    for para in doc.paragraphs:
        para.text = anonymize_text_transformer(para.text)
    doc.save(output_path)

@app.route('/anonymize', methods=['POST'])
def anonymize():
    """
    Handle file input and return two anonymized files (Presidio and Transformer).
    ---
    tags:
      - Anonymization
    consumes:
      - multipart/form-data
    parameters:
      - name: file
        in: formData
        type: file
        required: true
        description: The file to be anonymized (PDF or DOCX).
    responses:
      200:
        description: Zip file containing both anonymized versions.
        schema:
          type: file
      400:
        description: Invalid input or unsupported file type.
      500:
        description: Internal server error.
    """
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    uploaded_file = request.files['file']
    if uploaded_file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    file_extension = os.path.splitext(uploaded_file.filename)[1].lower()
    
    # Create temporary files
    input_temp = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
    presidio_output_temp = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
    transformer_output_temp = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
    zip_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.zip')
    
    input_path = input_temp.name
    presidio_output_path = presidio_output_temp.name
    transformer_output_path = transformer_output_temp.name
    zip_path = zip_temp.name

    try:
        uploaded_file.save(input_path)
        original_filename = os.path.basename(uploaded_file.filename)
        filename_without_ext = os.path.splitext(original_filename)[0]

        if file_extension == '.pdf':
            anonymize_pdf_presidio(input_path, presidio_output_path)
            anonymize_pdf_transformer(input_path, transformer_output_path)
        elif file_extension == '.docx':
            anonymize_docx_presidio(input_path, presidio_output_path)
            anonymize_docx_transformer(input_path, transformer_output_path)
        else:
            return jsonify({"error": "Unsupported file type"}), 400

        # Create a zip file containing both anonymized files
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            zipf.write(presidio_output_path, f"{filename_without_ext}_presidio{file_extension}")
            zipf.write(transformer_output_path, f"{filename_without_ext}_transformer{file_extension}")

        return send_file(
            zip_path, 
            as_attachment=True, 
            mimetype='application/zip', 
            download_name=f"anonymized_{filename_without_ext}.zip"
        )

    except Exception as e:
        logging.error(f"Error during anonymization: {str(e)}")
        return jsonify({"error": str(e)}), 500

    finally:
        # Clean up temporary files
        for temp_file in [input_path, presidio_output_path, transformer_output_path, zip_path]:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                logging.error(f"Error removing temporary file {temp_file}: {str(e)}")

@app.route('/')
def redirect_to_swagger():
    """Redirect to Swagger UI."""
    return redirect("/api-docs", code=302)


if __name__ == '__main__':
    logging.basicConfig(level=logging.DEBUG)
    app.run(host='0.0.0.0', port=5001, debug=True)